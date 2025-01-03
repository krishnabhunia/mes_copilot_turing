from docx import Document
import os
from pathlib import Path
# Create a new Word document
doc = Document()

# Add some content to the document
doc.add_heading('Document Title', level=1)
doc.add_paragraph('This is a sample paragraph.')

# Set document properties
core_properties = doc.core_properties
core_properties.title = "Sample Document"
core_properties.category = "Programming Documentation"
core_properties.subject = "Setting Metadata in Word"
core_properties.author = "Mr. Krishna"

# Save the document
os.makedirs("Output_Folder_Translated", exist_ok=True)
save_path = str(Path("Output_Folder_Translated/Sample_Document.docx"))
doc.save(save_path)
