import os
import shutil
from transformers import MarianMTModel, MarianTokenizer  # type: ignore
from docx import Document  # type: ignore
from PyPDF2 import PdfReader  # type: ignore

# Suppress TensorFlow logs
os.environ["TF_CPP_MIN_LOG_LEVEL"] = "2"
os.environ["TF_ENABLE_ONEDNN_OPTS"] = "0"

# Define paths
input_folder = 'input'
output_folder = 'translated_document'
os.environ["HF_HOME"] = "cache"

# Delete the folder if it exists
if os.path.exists(output_folder):
    print('Folder existed, so deleting ...')
    shutil.rmtree(output_folder)
    print('Folder deleted:', output_folder)

# Recreate the folder
os.makedirs(output_folder, exist_ok=True)

# Load MarianMT model and tokenizer
model_name = 'Helsinki-NLP/opus-mt-en-fr'
print("Tokenizer initializing ...")
tokenizer = MarianTokenizer.from_pretrained(model_name)
print("Tokenizer initialized, model initializing ...")
model = MarianMTModel.from_pretrained(model_name)
print("Model initialized")


# Translate function
def translate_text(text, tokenizer, model, chunk_size=400):
    print('Splitting text into chunks...')
    sentences = text.split("\n")
    translated = []
    current_chunk = []
    current_length = 0

    for sentence in sentences:
        if current_length + len(sentence) <= chunk_size:
            current_chunk.append(sentence)
            current_length += len(sentence)
        else:
            print('Translating chunk...')
            tokens = tokenizer.encode(" ".join(current_chunk), return_tensors='pt', max_length=512, truncation=True)
            translated_tokens = model.generate(tokens, max_length=512, num_beams=5, early_stopping=True)
            translated.append(tokenizer.decode(translated_tokens[0], skip_special_tokens=True))
            current_chunk = [sentence]
            current_length = len(sentence)

    if current_chunk:
        print('Translating remaining chunk...')
        tokens = tokenizer.encode(" ".join(current_chunk), return_tensors='pt', max_length=512, truncation=True)
        translated_tokens = model.generate(tokens, max_length=512, num_beams=5, early_stopping=True)
        translated.append(tokenizer.decode(translated_tokens[0], skip_special_tokens=True))

    return "\n".join(translated)


# Function to process and retain formatting in .docx
def translate_docx_with_formatting(input_path, output_path):
    doc = Document(input_path)
    translated_doc = Document()

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            print(f"Translating paragraph: {paragraph.text[:30]}...")  # Log first 30 characters
            translated_text = translate_text(paragraph.text, tokenizer, model)
        else:
            translated_text = ""

        # Add translated paragraph with the same style
        new_paragraph = translated_doc.add_paragraph(translated_text)
        new_paragraph.style = paragraph.style

    # Copy tables, headers, etc.
    for table in doc.tables:
        new_table = translated_doc.add_table(rows=len(table.rows), cols=len(table.columns))
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                translated_text = translate_text(cell.text, tokenizer, model) if cell.text.strip() else ""
                new_table.cell(i, j).text = translated_text

    translated_doc.save(output_path)


# Function to read .pdf file
def read_pdf(file_path):
    reader = PdfReader(file_path)
    return "\n".join([page.extract_text() for page in reader.pages])


# Function to get output file paths
def get_output_file_path(output_folder, file_name):
    output_file_name = f"translated_{file_name}"
    output_file_path = os.path.join(output_folder, output_file_name)
    return output_file_path, output_file_name


# Iterate through input folder
for file_name in os.listdir(input_folder):
    input_file_path = os.path.join(input_folder, file_name)

    if file_name.endswith('.docx'):
        print(f"Processing DOCX file: {file_name}")
        output_file_path, output_file_name = get_output_file_path(output_folder, file_name)
        translate_docx_with_formatting(input_file_path, output_file_path)
        print(f"Translated DOCX saved: {output_file_name}")

    elif file_name.endswith('.pdf'):
        print(f"Processing PDF file: {file_name}")
        content = read_pdf(input_file_path)
        output_file_path, output_file_name = get_output_file_path(output_folder, file_name)
        translated_content = translate_text(content, tokenizer, model)
        with open(output_file_path, 'w', encoding='utf-8') as f:
            f.write(translated_content)
        print(f"Translated PDF saved as text: {output_file_name}")

    elif file_name.endswith('.txt'):
        print(f"Processing TXT file: {file_name}")
        with open(input_file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        translated_content = translate_text(content, tokenizer, model)
        output_file_path, output_file_name = get_output_file_path(output_folder, file_name)
        with open(output_file_path, 'w', encoding='utf-8') as f:
            f.write(translated_content)
        print(f"Translated TXT saved: {output_file_name}")

print("Translation completed.")
