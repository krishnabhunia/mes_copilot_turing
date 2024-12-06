from docx import Document  # type: ignore
from transformers import MarianMTModel, MarianTokenizer  # type: ignore

# File paths
input_file_path = '/mnt/data/Krishna.docx'
output_file_path = '/mnt/data/Translated_Krishna.docx'

# Load MarianMT model and tokenizer
model_name = 'Helsinki-NLP/opus-mt-en-fr'
tokenizer = MarianTokenizer.from_pretrained(model_name)
model = MarianMTModel.from_pretrained(model_name)


# Translate text in chunks
def translate_text(text, tokenizer, model, chunk_size=400):
    sentences = text.split("\n")
    translated = []
    current_chunk = []
    current_length = 0

    for sentence in sentences:
        if current_length + len(sentence) <= chunk_size:
            current_chunk.append(sentence)
            current_length += len(sentence)
        else:
            tokens = tokenizer.encode(" ".join(current_chunk), return_tensors='pt', max_length=512, truncation=True)
            translated_tokens = model.generate(tokens, max_length=512, num_beams=5, early_stopping=True)
            translated.append(tokenizer.decode(translated_tokens[0], skip_special_tokens=True))
            current_chunk = [sentence]
            current_length = len(sentence)

    if current_chunk:
        tokens = tokenizer.encode(" ".join(current_chunk), return_tensors='pt', max_length=512, truncation=True)
        translated_tokens = model.generate(tokens, max_length=512, num_beams=5, early_stopping=True)
        translated.append(tokenizer.decode(translated_tokens[0], skip_special_tokens=True))

    return "\n".join(translated)


# Translate document while retaining formatting
def translate_docx_with_formatting(input_path, output_path):
    doc = Document(input_path)
    translated_doc = Document()

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            translated_text = translate_text(paragraph.text, tokenizer, model)
        else:
            translated_text = ""

        new_paragraph = translated_doc.add_paragraph(translated_text)
        new_paragraph.style = paragraph.style
        new_paragraph.alignment = paragraph.alignment
        new_paragraph.paragraph_format.left_indent = paragraph.paragraph_format.left_indent
        new_paragraph.paragraph_format.right_indent = paragraph.paragraph_format.right_indent
        new_paragraph.paragraph_format.first_line_indent = paragraph.paragraph_format.first_line_indent
        new_paragraph.paragraph_format.line_spacing = paragraph.paragraph_format.line_spacing

    # Translate tables
    for table in doc.tables:
        new_table = translated_doc.add_table(rows=len(table.rows), cols=len(table.columns))
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                translated_text = translate_text(cell.text, tokenizer, model) if cell.text.strip() else ""
                new_table.cell(i, j).text = translated_text

    translated_doc.save(output_path)


# Perform translation
translate_docx_with_formatting(input_file_path, output_file_path)
