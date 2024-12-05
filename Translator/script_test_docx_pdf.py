import os
import shutil
from transformers import MarianMTModel, MarianTokenizer  # type: ignore
from docx import Document  # For .docx files   # type: ignore
from PyPDF2 import PdfReader  # For .pdf files   # type: ignore

# Suppress TensorFlow logs
os.environ["TF_CPP_MIN_LOG_LEVEL"] = "2"
os.environ["TF_ENABLE_ONEDNN_OPTS"] = "0"

# Define paths
input_folder = 'input'
output_folder = 'translated_document'
os.environ["HF_HOME"] = "cache"


# Delete the folder if it exists
if os.path.exists(output_folder):
    print('folder existed, so deleting ...')
    shutil.rmtree(output_folder)
    print('folder deleted : ', output_folder)

# Recreate the folder
os.makedirs(output_folder, exist_ok=True)

# Load MarianMT model and tokenizer
model_name = 'Helsinki-NLP/opus-mt-en-fr'
print("tokenizer intializing ...")
tokenizer = MarianTokenizer.from_pretrained(model_name)
print("tokenizer initialized, model initializing ...")
model = MarianMTModel.from_pretrained(model_name)
print("model initialized")


# Translate function
def translate_text(text, tokenizer, model):
    print('token encoding ...')
    tokens = tokenizer.encode(text, return_tensors='pt', max_length=512, truncation=True)
    print('token encoded, model generating tokens ...')
    translated_tokens = model.generate(tokens, max_length=512, num_beams=5, early_stopping=True)
    print('translated token generated, returning tokens')
    return tokenizer.decode(translated_tokens[0], skip_special_tokens=True)


# Function to read .docx file
def read_docx(file_path):
    doc = Document(file_path)
    return "\n".join([paragraph.text for paragraph in doc.paragraphs])


# Function to read .pdf file
def read_pdf(file_path):
    reader = PdfReader(file_path)
    return "\n".join([page.extract_text() for page in reader.pages])


# Function to write translated content to .docx file
def write_docx(file_path, content):
    doc = Document()
    for line in content.split("\n"):
        doc.add_paragraph(line)
    doc.save(file_path)


def get_output_file_path(output_folder, file_name):
    output_file_name = f"translated_{file_name}"
    output_file_path = os.path.join(output_folder, output_file_name)
    return output_file_path, output_file_name


# Iterate through input folder
for file_name in os.listdir(input_folder):
    input_file_path = os.path.join(input_folder, file_name)

    if file_name.endswith('.docx'):
        # Process .docx file
        print(f"Processing DOCX file: {file_name}")
        content = read_docx(input_file_path)
        translated_content = translate_text(content, tokenizer, model)
        output_file_path, output_file_name = get_output_file_path(output_folder, file_name)
        write_docx(output_file_path, translated_content)
        print(f"Translated DOCX saved: {output_file_name}")

    elif file_name.endswith('.pdf'):
        # Process .pdf file
        print(f"Processing PDF file: {file_name}")
        content = read_pdf(input_file_path)
        output_file_path, output_file_name = get_output_file_path(output_folder, file_name)
        translated_content = translate_text(content, tokenizer, model)
        with open(output_file_path, 'w', encoding='utf-8') as f:
            f.write(translated_content)
        print(f"Translated PDF saved as text: {output_file_name}")

    elif file_name.endswith('.txt'):
        # Process .txt file
        print(f"Processing TXT file: {file_name}")
        with open(input_file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        translated_content = translate_text(content, tokenizer, model)
        output_file_path, output_file_name = get_output_file_path(output_folder, file_name)
        with open(output_file_path, 'w', encoding='utf-8') as f:
            f.write(translated_content)
        print(f"Translated TXT saved: {output_file_name}")

print("Translation completed.")
